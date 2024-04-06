from docx import Document

# 打开Word文件
doc = Document('饼干检验报告（第二份）.docx')

# 遍历文档中的段落
for paragraph in doc.paragraphs:
    # 在这里可以添加条件判断，选择要修改的段落
    # 这里演示将所有段落中包含 'old content' 的部分替换为 'New Content'
    print(paragraph.text)
    # if 'old content' in paragraph.text:
    #     # 修改段落的文本
    #     paragraph.text = paragraph.text.replace('old content', 'New Content')

# 遍历文档中的表格
for table in doc.tables:
    # 遍历表格中的单元格
    for row in table.rows:
        print('row:', row)
        for cell in row.cells:
            print(cell.text)
            if cell.text == '':
                cell.text = '1'
            # 在这里可以添加条件判断，选择要修改的单元格
            # 这里演示将所有单元格中包含 'old content' 的部分替换为 'New Content'
            # if 'old content' in cell.text:
            #     # 修改单元格的文本
            #     cell.text = cell.text.replace('old content', 'New Content')

# 保存修改后的Word文件
doc.save('修改后饼干检验报告（第二份）.docx')
